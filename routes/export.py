import io
import os
from datetime import datetime
from flask import Blueprint, send_file, jsonify

from models.meeting import Meeting
from middleware.auth import auth_required

export_bp = Blueprint("export", __name__)


# ── GET /api/export/<id>/pdf ──────────────────────────────────────────────────
@export_bp.route("/<meeting_id>/pdf", methods=["GET"])
@auth_required
def export_pdf(current_user, meeting_id):
    meeting = Meeting.query.filter_by(id=meeting_id, user_id=current_user.id).first()
    if not meeting:
        return jsonify({"success": False, "message": "Meeting not found"}), 404

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        HRFlowable, PageBreak, ListItem, ListFlowable)
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError:
        return jsonify({"success": False, "message": "reportlab not installed"}), 500

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    accent = colors.HexColor("#6366f1")
    dark   = colors.HexColor("#1a1a2e")
    grey   = colors.HexColor("#6b7280")

    title_style = ParagraphStyle("Title", parent=styles["Title"],
                                 fontSize=22, textColor=dark, alignment=TA_CENTER, spaceAfter=4)
    h1_style    = ParagraphStyle("H1", parent=styles["Heading1"],
                                 fontSize=16, textColor=dark, spaceBefore=12, spaceAfter=6)
    h2_style    = ParagraphStyle("H2", parent=styles["Heading2"],
                                 fontSize=13, textColor=accent, spaceBefore=10, spaceAfter=4)
    body_style  = ParagraphStyle("Body", parent=styles["Normal"],
                                 fontSize=11, textColor=colors.HexColor("#374151"), leading=16)
    meta_style  = ParagraphStyle("Meta", parent=styles["Normal"],
                                 fontSize=10, textColor=grey, alignment=TA_CENTER)
    bullet_style = ParagraphStyle("Bullet", parent=body_style, leftIndent=14, spaceAfter=3)

    story = []

    # Header
    story.append(Paragraph("Meeting Notes", title_style))
    story.append(Paragraph(meeting.title, h1_style))
    duration_min = round((meeting.duration or 0) / 60)
    story.append(Paragraph(
        f"Date: {meeting.created_at.strftime('%B %d, %Y')} &nbsp;&nbsp;|&nbsp;&nbsp; "
        f"Duration: {duration_min} min &nbsp;&nbsp;|&nbsp;&nbsp; Words: {meeting.word_count or 0}",
        meta_style,
    ))
    story.append(Spacer(1, 10))
    story.append(HRFlowable(width="100%", color=colors.HexColor("#e2e8f0")))
    story.append(Spacer(1, 10))

    # Summary
    summary = meeting.summary or {}
    if summary.get("overview"):
        story.append(Paragraph("Meeting Summary", h2_style))
        story.append(Paragraph(summary["overview"], body_style))
        story.append(Spacer(1, 8))

    if summary.get("keyPoints"):
        story.append(Paragraph("Key Discussion Points", h2_style))
        items = [ListItem(Paragraph(f"• {p}", bullet_style)) for p in summary["keyPoints"]]
        story.append(ListFlowable(items, bulletType="bullet", leftIndent=10))
        story.append(Spacer(1, 8))

    # Action items
    action_items = meeting.action_items or []
    if action_items:
        story.append(Paragraph("Action Items", h2_style))
        for i, item in enumerate(action_items, 1):
            status   = "☑" if item.get("completed") else "☐"
            assignee = f" → {item['assignee']}" if item.get("assignee") else ""
            priority = f"[{item.get('priority', 'medium').upper()}]"
            story.append(Paragraph(f"{status} {i}. {item['task']}{assignee} {priority}", bullet_style))
        story.append(Spacer(1, 8))

    # Participants
    participants = meeting.participants or []
    if participants:
        story.append(Paragraph("Participants", h2_style))
        for p in participants:
            role = f" ({p['role']})" if p.get("role") else ""
            story.append(Paragraph(f"• {p['name']}{role}", bullet_style))
        story.append(Spacer(1, 8))

    # Transcript on new page
    transcript = meeting.transcript or {}
    if transcript.get("full"):
        story.append(PageBreak())
        story.append(Paragraph("Full Transcript", h2_style))
        story.append(Spacer(1, 6))
        segments = transcript.get("segments", [])
        if segments:
            for seg in segments:
                story.append(Paragraph(f"<b>{seg.get('speaker', 'Speaker')}:</b>", body_style))
                story.append(Paragraph(seg.get("text", ""), body_style))
                story.append(Spacer(1, 4))
        else:
            story.append(Paragraph(transcript["full"], body_style))

    doc.build(story)
    buf.seek(0)

    safe_name = "".join(c if c.isalnum() else "_" for c in meeting.title)
    return send_file(buf, mimetype="application/pdf",
                     as_attachment=True, download_name=f"{safe_name}.pdf")


# ── GET /api/export/<id>/docx ─────────────────────────────────────────────────
@export_bp.route("/<meeting_id>/docx", methods=["GET"])
@auth_required
def export_docx(current_user, meeting_id):
    meeting = Meeting.query.filter_by(id=meeting_id, user_id=current_user.id).first()
    if not meeting:
        return jsonify({"success": False, "message": "Meeting not found"}), 404

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return jsonify({"success": False, "message": "python-docx not installed"}), 500

    doc = Document()

    # Style helpers
    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        return h

    def add_para(text, bold=False, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        run.font.size = Pt(11)
        return p

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Meeting Notes")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    add_heading(meeting.title, level=1)

    duration_min = round((meeting.duration or 0) / 60)
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.add_run(
        f"Date: {meeting.created_at.strftime('%B %d, %Y')}  |  "
        f"Duration: {duration_min} min  |  Words: {meeting.word_count or 0}"
    ).font.size = Pt(10)

    doc.add_paragraph()

    # Summary
    summary = meeting.summary or {}
    if summary.get("overview"):
        add_heading("Meeting Summary", level=2)
        add_para(summary["overview"])

    if summary.get("keyPoints"):
        add_heading("Key Discussion Points", level=2)
        for point in summary["keyPoints"]:
            doc.add_paragraph(f"• {point}", style="List Bullet")

    # Action items
    action_items = meeting.action_items or []
    if action_items:
        add_heading("Action Items", level=2)
        for i, item in enumerate(action_items, 1):
            status   = "☑" if item.get("completed") else "☐"
            assignee = f" → {item['assignee']}" if item.get("assignee") else ""
            priority = f" [{item.get('priority', 'medium').upper()}]"
            doc.add_paragraph(f"{status} {i}. {item['task']}{assignee}{priority}")

    # Participants
    participants = meeting.participants or []
    if participants:
        add_heading("Participants", level=2)
        for p in participants:
            role = f" ({p['role']})" if p.get("role") else ""
            doc.add_paragraph(f"• {p['name']}{role}")

    # Transcript
    transcript = meeting.transcript or {}
    if transcript.get("full"):
        doc.add_page_break()
        add_heading("Full Transcript", level=2)
        segments = transcript.get("segments", [])
        if segments:
            for seg in segments:
                p = doc.add_paragraph()
                r = p.add_run(f"{seg.get('speaker', 'Speaker')}: ")
                r.bold = True
                p.add_run(seg.get("text", ""))
        else:
            add_para(transcript["full"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = "".join(c if c.isalnum() else "_" for c in meeting.title)
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"{safe_name}.docx",
    )
